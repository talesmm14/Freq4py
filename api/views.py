from os import name, system
from django.contrib.auth.models import User
from django.shortcuts import render
from django.http import Http404, StreamingHttpResponse
from django.conf import settings
from django.shortcuts import get_object_or_404, render, redirect
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.utils import serializer_helpers
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import generics, status
from api.models import Not_Work_Type, Not_Working_Day, Schedule, Sheet, Sheet_Title, Sheet_Value
from api.serializers import Not_Working_Type_Serializer, Schedule_Serializer, Sheet_Serializer, RegisterSerializer, \
    Titles_fields_Serializer, Values_fields_Serializer

import io
import sys
import docx
import datetime
import locale
import os
from calendar import monthrange

locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')

from django.contrib.auth import get_user_model

User = get_user_model()


class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    permission_classes = (AllowAny,)
    serializer_class = RegisterSerializer


class Not_Work_Type_List(generics.ListCreateAPIView):
    queryset = Not_Work_Type.objects.all()
    serializer_class = Not_Working_Type_Serializer


class Schedule_List(generics.ListCreateAPIView):
    queryset = Schedule.objects.all()
    serializer_class = Schedule_Serializer


class Schedule_Detail(generics.RetrieveUpdateDestroyAPIView):
    queryset = Schedule.objects.all()
    serializer_class = Schedule_Serializer


class Titles_List(generics.ListCreateAPIView):
    queryset = Sheet_Title.objects.all()
    serializer_class = Titles_fields_Serializer


class Titles_Detail(generics.RetrieveUpdateDestroyAPIView):
    queryset = Sheet_Title.objects.all()
    serializer_class = Titles_fields_Serializer


class Values_List(APIView):
    def get(self, request, format=None):
        snippets = Sheet_Value.objects.all()
        serializer = Values_fields_Serializer(snippets, many=True)
        return Response(serializer.data)

    def post(self, request, format=None):
        serializer = Values_fields_Serializer(
            data=request.data, context={'request': request})
        if serializer.is_valid():
            value = serializer.save()
            value.user = self.request.user
            value.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class Values_Detail(APIView):
    def get_object(self, pk, user):
        try:
            value = Sheet_Value.objects.get(pk=pk)
            if value.user == user:
                return value
            return None
        except Sheet_Value.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        snippets = self.get_object(pk, self.request.user)
        serializer = Values_fields_Serializer(snippets, many=True)
        return Response(serializer.data)

    def put(self, request, pk, format=None):
        snippets = self.get_object(pk, self.request.user)
        serializer = Values_fields_Serializer(snippets, many=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk, format=None):
        snippets = self.get_object(pk, self.request.user)
        serializer = Values_fields_Serializer(snippets, many=True)

        if serializer.is_valid():
            serializer.delete()
            return Response(status.HTTP_200_OK)
        return Response(serializer.errors, status.HTTP_400_BAD_REQUEST)


class Sheet_List(APIView):
    def get(self, request, format=None):
        snippets = Sheet.objects.all()
        serializer = Sheet_Serializer(snippets, many=True)
        return Response(serializer.data)

    def post(self, request, format=None):
        serializer = Sheet_Serializer(
            data=request.data, context={'request': request})
        if serializer.is_valid():
            sheet = serializer.save()
            sheet.user = self.request.user
            sheet.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class Sheet_Detail(APIView):
    def get_object(self, pk):
        try:
            sheet = Sheet.objects.get(pk=pk)
            return sheet
        except Sheet.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = Sheet_Serializer(snippets, many=True)
        return Response(serializer.data)

    def put(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = Sheet_Serializer(snippets, many=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = Sheet_Serializer(snippets, many=True)

        if serializer.is_valid():
            serializer.delete()
            return Response(status.HTTP_200_OK)
        return Response(serializer.errors, status.HTTP_400_BAD_REQUEST)


class ExportDocx(APIView):
    def get_object(self, pk):
        try:
            sheet = Sheet.objects.get(pk=pk)
            return sheet
        except Sheet.DoesNotExist:
            raise Http404

    def get(self, request, pk, *args, **kwargs):
        sheet = self.get_object(pk)
        # create an empty document object
        # print(settings.STATIC_ROOT + sheet.path, file=sys.stderr)

        path = settings.STATIC_ROOT + sheet.path
        key_words = get_keys(sheet)

        if key_words == None:
            return Response(status.HTTP_400_BAD_REQUEST)

        document = docx.Document(path)

        replace(document, sheet, key_words)

        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream

        # put them to streaming content response
        # within docx content_type
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment;filename={}.docx'.format(sheet.values_fields.field_value_1)
        response["Content-Encoding"] = 'UTF-8'

        return response


def get_keys(sheet):
    if (sheet.titles_fields != None):
        title_fields = sheet.titles_fields
        if (sheet.values_fields == None):
            value_fields = Sheet_Value()
        else:
            value_fields = sheet.values_fields
        key_words = {**title_fields.__dict__, **value_fields.__dict__}
    else:
        return None

    key_words.pop('_state')
    key_words.pop('id')
    key_words.pop('name')
    return key_words


def replace(doc, sheet, key_words):
    year = sheet.date.year
    month = sheet.date.month
    month_days = monthrange(year, month)[1]
    key_table = sheet.schedule.key_words()
    not_working_days = {int(not_work.day): str(not_work.description)
                        for not_work in Not_Working_Day.objects.filter(sheet=sheet.id)}
    key_table.update({'field_date': (datetime.date(
        year, month, 1).strftime('%B')).swapcase() + "/" + str(year)})

    key_words = {**key_words, **key_table}
    day = 0
    weekday = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == "HM1":
                    day += 1
                    if day <= month_days:
                        weekday = datetime.date(
                            year=year, month=month, day=day).weekday()

                for paragraph in cell.paragraphs:
                    for variable_key, variable_value in key_words.items():
                        if weekday == 5 or weekday == 6:
                            variable_value = replace_weekend(
                                weekday, variable_key)
                        elif day in not_working_days.keys():
                            if variable_key[0] != 'H':
                                variable_value = not_working_days[day]
                            else:
                                variable_value = "****"
                        if variable_value == 'HM1':
                            print(paragraph.text, variable_key, variable_value)
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value)


def replace_weekend(weekday, key):
    if key[0] != 'H':
        if weekday == 5:
            return "SÁBADO"
        if weekday == 6:
            return "DOMINGO"
    return "****"


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                if value != None:
                    item.text = item.text.replace(key, value)
                else:
                    item.text = ""
