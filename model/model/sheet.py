import docx
from calendar import monthrange
import datetime
from datetime import date
import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')


class Sheet():
    def __init__(self, path, save_path, year, month, key_words, not_working_days=None):
        self.doc = docx.Document(path)
        self.path = path
        self.save_path = save_path
        self.year = year
        self.month = month
        self.key_words = key_words
        self.not_working_days = not_working_days
        self.month_days = monthrange(year, month)[1]
        self.year_month_title = (datetime.date(year, month, 1).strftime('%B')).swapcase() + "/" + str(year)
        self.key_table = {
            "AM1": "",
            "AM2": "",
            "AV1": "",
            "AV2": "",
            "HM1": "08:00",
            "HM2": "14:00",
            "HV1": "12:00",
            "HV2": "18:00",
        }

    def replace_key_table(self, key_table):
        self.key_table = key_table

    def __replace(self):
        self.key_words["${FIELD_DATE}"] = str(datetime.date(self.year, self.month, 1).strftime('%B')).swapcase() + "/" + str(self.year)
        day = 0
        weekday = 0
        for variable_key, variable_value in self.key_words.items():
            for paragraph in self.doc.paragraphs:
                self.__replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text == "HM1":
                            day += 1
                            if day <= self.month_days:
                                weekday = date(
                                    year=self.year, month=self.month, day=day).weekday()

                        for key, value in self.key_table.items():
                            value = self.__not_working_days(day, weekday, key, value)
                            self.__replace_text_in_paragraph(paragraph, key, value)

                        for paragraph in cell.paragraphs:
                            self.__replace_text_in_paragraph(paragraph, variable_key, variable_value)

    def __not_working_days(self, weekday, day, key, value):
        if weekday == 5 or weekday == 6:
            if key[0] != 'H':
                if weekday == 5:
                    value = "SÁBADO"
                if weekday == 6:
                    value = "DOMINGO"
            else:
                value = "****"
        elif self.not_working_days != None:
            if key[0] != 'H':
                for not_working_day, not_working_day_value in self.not_working_days.items():
                    if day == not_working_day:
                        value = not_working_day_value
        return value

    def __replace_text_in_paragraph(self, paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    def save(self):
        self.__replace()
        self.doc.save(str(self.save_path + '/Frequência_{}.docx'.format(self.key_words['${FIELD_VALUE_1}'])))
