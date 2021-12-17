from django.http import Http404, StreamingHttpResponse
from django.conf import settings
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import generics, status
from django.contrib.auth import get_user_model

from api.models import NotWorkType, NotWorkingDay, Schedule, Sheet, SheetTitle, SheetValue
from api.permissions import ReadOnly
from api.serializers import NotWorkingTypeSerializer, ScheduleSerializer, SheetSerializer, RegisterSerializer, \
    TitlesFieldsSerializer, ValuesFieldsSerializer, NotWorkingDaySerializer

import io
import docx
import datetime
import locale
from calendar import monthrange

locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')

User = get_user_model()


class RegisterView(generics.CreateAPIView):
    queryset = User.objects.all()
    permission_classes = (AllowAny,)
    serializer_class = RegisterSerializer


class NotWorkTypeList(generics.ListCreateAPIView):
    queryset = NotWorkType.objects.all()
    serializer_class = NotWorkingTypeSerializer
    permission_classes = [IsAuthenticated]


class NotWorkTypeDetail(generics.RetrieveUpdateDestroyAPIView):
    queryset = NotWorkType.objects.all()
    serializer_class = NotWorkingTypeSerializer
    permission_classes = [IsAuthenticated, ReadOnly]


class ScheduleList(generics.ListCreateAPIView):
    queryset = Schedule.objects.all()
    serializer_class = ScheduleSerializer
    permission_classes = [IsAuthenticated]


class ScheduleDetail(generics.RetrieveUpdateDestroyAPIView):
    queryset = Schedule.objects.all()
    serializer_class = ScheduleSerializer
    permission_classes = [IsAuthenticated, ReadOnly]


class NotWorkingDayList(generics.ListCreateAPIView):
    queryset = NotWorkingDay.objects.all()
    serializer_class = NotWorkingDaySerializer
    permission_classes = [IsAuthenticated]


class NotWorkingDayDetail(generics.RetrieveUpdateDestroyAPIView):
    queryset = NotWorkingDay.objects.all()
    serializer_class = NotWorkingDaySerializer
    permission_classes = [IsAuthenticated, ReadOnly]


class TitlesList(generics.ListCreateAPIView):
    queryset = SheetTitle.objects.all()
    serializer_class = TitlesFieldsSerializer
    permission_classes = [IsAuthenticated]


class TitlesDetail(generics.RetrieveUpdateDestroyAPIView):
    queryset = SheetTitle.objects.all()
    serializer_class = TitlesFieldsSerializer
    permission_classes = [IsAuthenticated, ReadOnly]


class ValuesList(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, format=None):
        snippets = User.objects.get(pk=request.user.pk).sheetvalue_set.all()
        serializer = ValuesFieldsSerializer(snippets, many=True)
        return Response(serializer.data)

    def post(self, request, format=None):
        serializer = ValuesFieldsSerializer(
            data=request.data, context={'request': request})
        if serializer.is_valid():
            value = serializer.save()
            value.user = self.request.user
            value.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ValuesDetail(APIView):
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        try:
            return User.objects.get(pk=self.request.user.pk).sheetvalue_set.filter(pk=pk)
        except SheetValue.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = ValuesFieldsSerializer(snippets, many=True)
        return Response(serializer.data)

    def put(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = ValuesFieldsSerializer(snippets, many=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk, format=None):
        snippets = self.get_object(pk)

        if snippets.is_valid():
            snippets.delete()
            return Response(status.HTTP_200_OK)
        return Response(status.HTTP_400_BAD_REQUEST)


class SheetList(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, format=None):
        snippets = User.objects.get(pk=self.request.user.pk).sheet_set.all()
        serializer = SheetSerializer(snippets, many=True)
        return Response(serializer.data)

    def post(self, request, format=None):
        serializer = SheetSerializer(
            data=request.data, context={'request': request})
        if serializer.is_valid():
            sheet = serializer.save()
            sheet.user = self.request.user
            sheet.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class SheetDetail(APIView):
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        try:
            sheet = User.objects.get(pk=self.request.user.pk).sheet_set.filter(pk=pk)
            return sheet
        except Sheet.DoesNotExist:
            raise Http404

    def get(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = SheetSerializer(snippets, many=True)
        return Response(serializer.data)

    def put(self, request, pk, format=None):
        snippets = self.get_object(pk)
        serializer = SheetSerializer(snippets, many=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk, format=None):
        snippets = self.get_object(pk)

        if snippets.is_valid():
            snippets.delete()
            return Response(status.HTTP_200_OK)
        return Response(status.HTTP_400_BAD_REQUEST)


class ExportDocx(APIView):
    permission_classes = [IsAuthenticated]

    def get_object(self, pk):
        try:
            sheet = User.objects.get(pk=self.request.user.pk).sheet_set.filter(pk=pk)
            return sheet
        except Sheet.DoesNotExist:
            raise Http404

    def get(self, request, pk, *args, **kwargs):
        sheet = self.get_object(pk)
        # create an empty document object
        # print(settings.STATIC_ROOT + sheet.path, file=sys.stderr)

        path = settings.STATIC_ROOT + sheet.path
        key_words = get_keys(sheet)

        if key_words is None:
            return Response(status.HTTP_400_BAD_REQUEST)

        document = docx.Document(path)

        replace(document, sheet, key_words)

        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream

        # put them to streaming content response
        # within docx content_type
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment;filename={}.docx'.format(sheet.values_fields.field_value_1)
        response["Content-Encoding"] = 'UTF-8'

        return response


def get_keys(sheet):
    if sheet.titles_fields is not None:
        title_fields = sheet.titles_fields
        if sheet.values_fields is None:
            value_fields = SheetValue()
        else:
            value_fields = sheet.values_fields
        key_words = {**title_fields.__dict__, **value_fields.__dict__}
    else:
        return None

    key_words.pop('_state')
    key_words.pop('id')
    key_words.pop('name')
    return key_words


def replace(doc, sheet, key_words):
    year = sheet.date.year
    month = sheet.date.month
    month_days = monthrange(year, month)[1]
    key_table = sheet.schedule.key_words()
    not_working_days = {int(not_work.day): str(not_work.description)
                        for not_work in NotWorkingDay.objects.filter(sheet=sheet.id)}
    key_table.update({'field_date': (datetime.date(
        year, month, 1).strftime('%B')).swapcase() + "/" + str(year)})

    key_words = {**key_words, **key_table}
    day = 0
    weekday = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == "HM1":
                    day += 1
                    if day <= month_days:
                        weekday = datetime.date(
                            year=year, month=month, day=day).weekday()

                for paragraph in cell.paragraphs:
                    for variable_key, variable_value in key_words.items():
                        if weekday == 5 or weekday == 6:
                            variable_value = replace_weekend(
                                weekday, variable_key)
                        elif day in not_working_days.keys():
                            if variable_key[0] != 'H':
                                variable_value = not_working_days[day]
                            else:
                                variable_value = "****"
                        if variable_value == 'HM1':
                            print(paragraph.text, variable_key, variable_value)
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value)


def replace_weekend(weekday, key):
    if key[0] != 'H':
        if weekday == 5:
            return "SÁBADO"
        if weekday == 6:
            return "DOMINGO"
    return "****"


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                if value is not None:
                    item.text = item.text.replace(key, value)
                else:
                    item.text = ""
