import docx
from calendar import month, monthrange, weekday
import datetime
from datetime import date
import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')

ano = 2021
mes = 9
mes_max_dia = monthrange(ano, mes)[1]

key_table = {
    "AS": "",
    "HM1": "08:00",
    "HM2": "14:00",
    "HV1": "12:00",
    "HV2": "18:00",
}

def cargoComissionado(key_words):
    key_words["${FIELD_TITLE_1}"] = "NOME:"
    key_words["${FIELD_TITLE_2}"] =  "LOTAÇÃO:"
    key_words["${FIELD_TITLE_3}"] =  "CARGO EFETIVO:"
    key_words["${FIELD_TITLE_4}"] =  "CARGO COMISSIONADO:"
    key_words["${FIELD_TITLE_5}"] =  "MATRÍCULA:"
    key_words["${FIELD_TITLE_6}"] =  "HORÁRIO:"
    key_words["${FIELD_TITLE_7}"] =  "Chefe Imediato:"
    key_words["${FIELD_TITLE_8}"] =  "Chefe Mediato:"

def cargoEstagiario(key_words):
    key_words["${FIELD_TITLE_1}"] = "NOME:"
    key_words["${FIELD_TITLE_2}"] =  "SUPERVISOR:"
    key_words["${FIELD_TITLE_3}"] =  "CARGO:"
    key_words["${FIELD_TITLE_4}"] =  ""
    key_words["${FIELD_TITLE_5}"] =  ""
    key_words["${FIELD_TITLE_6}"] =  "HORÁRIO:"
    key_words["${FIELD_TITLE_7}"] =  "SUPERVISOR:"
    key_words["${FIELD_TITLE_8}"] =  "MATRÍCULA:"

def main():
    path = '../Freq4py/Frequência - Modelo.docx'

    doc = docx.Document(path)

    key_words = {
        "${FIELD_DATE}": (datetime.date(ano, mes, 1).strftime('%B')).swapcase() + "/" + str(ano),
        "${FIELD_TITLE_1}": "",
        "${FIELD_TITLE_2}": "",
        "${FIELD_TITLE_3}": "",
        "${FIELD_TITLE_4}": "",
        "${FIELD_TITLE_5}": "",
        "${FIELD_TITLE_6}": "",
        "${FIELD_TITLE_7}": "",
        "${FIELD_TITLE_8}": "",

        "${FIELD_VALUE_1}": "Tales Monteiro Melquiades",
        "${FIELD_VALUE_2}": "Teste Lotacao",
        "${FIELD_VALUE_3}": "Estagiario",
        "${FIELD_VALUE_4}": "Não",
        "${FIELD_VALUE_5}": "123453454",
        "${FIELD_VALUE_6}": "12:00 - 18:00",
        "${FIELD_VALUE_7}": "Stefan",
        "${FIELD_VALUE_8}": "Stefan",
    }

    cargoEstagiario(key_words)

    replace_table_values(doc.tables)

    for variable_key, variable_value in key_words.items():
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value)

    doc.save(
        '../Freq4py/save_doc/Frequência_{}.docx'.format(key_words['${FIELD_VALUE_1}']))


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def replace_table_values(tables):
    day = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == "HM1":
                    for paragraph in cell.paragraphs:
                        day+=1
                        if day <= mes_max_dia:
                            weekday = date(year=ano, month=mes, day=day).weekday()
                        paragraph.text = key_table["HM1"]
                        
                if cell.text == "AM1":
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph_in_table(paragraph, "AS", weekday)

                if cell.text == "HM2":
                    for paragraph in cell.paragraphs:
                        paragraph.text = key_table["HM2"]

                if cell.text == "AM2":
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph_in_table(paragraph, "AS", weekday)

                if cell.text == "HV1":
                    for paragraph in cell.paragraphs:
                        paragraph.text = key_table["HV1"]

                if cell.text == "AV1":
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph_in_table(paragraph, "AS", weekday)

                if cell.text == "HV2":
                    for paragraph in cell.paragraphs:
                        paragraph.text = key_table["HV2"]

                if cell.text == "AV2":
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph_in_table(paragraph, "AS", weekday)

def replace_text_in_paragraph_in_table(paragraph, key, weekday):
    if weekday == 5:
        paragraph.text = "SÁBADO"
        return

    if weekday == 6:
        paragraph.text = "DOMINGO"
        return

    paragraph.text = key_table[key]


if __name__ == '__main__':
    main()
