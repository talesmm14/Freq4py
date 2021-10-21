import docx
from calendar import month, monthrange, weekday
import datetime
from datetime import date
import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')


class Sheet():
    def __init__(self, path, save_path, year, month, key_words):
        self.doc = docx.Document(path)
        self.path = path
        self.save_path = save_path
        self.year = year
        self.month = month
        self.key_words = key_words
        self.month_days = monthrange(year, month)[1]
        self.year_month_title = (datetime.date(
            year, month, 1).strftime('%B')).swapcase() + "/" + str(year)
        self.key_table = {
            "AM1": "",
            "AM2": "",
            "AV1": "",
            "AV2": "",
            "HM1": "08:00",
            "HM2": "14:00",
            "HV1": "12:00",
            "HV2": "18:00",
        }

    def replace_key_table(self, key_table):
        self.key_table = key_table

    def __replace(self):
        day = 0
        weekday = 0
        for variable_key, variable_value in self.key_words.items():
            for paragraph in self.doc.paragraphs:
                self.__replace_text_in_paragraph(
                    paragraph, variable_key, variable_value)

            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text == "HM1":
                            day += 1
                            if day <= self.month_days:
                                weekday = date(
                                    year=self.year, month=self.month, day=day).weekday()

                        for key, value in self.key_table.items():
                            if weekday > 5 and weekday < 6:
                                value = self.__weekend_value()
                            self.__replace_text_in_paragraph(
                                paragraph, key, value)

                        for paragraph in cell.paragraphs:
                            self.__replace_text_in_paragraph(
                                paragraph, variable_key, variable_value)

    def __weekend_value(weekday, key):
        if key[0] != 'H':
            if weekday == 5:
                return "SÁBADO"
            if weekday == 6:
                return "DOMINGO"
        return "..."

    def __replace_text_in_paragraph(paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)

    def save(self):
        self.__replace()
        self.doc.save(
            self.save_path + '/Frequência_{}.docx'.format(key_words['${FIELD_VALUE_1}']))
